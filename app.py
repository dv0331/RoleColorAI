"""
RoleColorAI - Main Streamlit Application

A resume intelligence system that analyzes and rewrites resumes
through a team-role lens (Builder, Enabler, Thriver, Supportee).

Features:
- RoleColor scoring and analysis
- AI-powered summary rewriting
- LaTeX/Overleaf code generation
- E2B sandbox for LaTeX compilation
- Interactive chat assistant
- PDF and Word document upload support
"""

import streamlit as st
import base64
import subprocess
import tempfile
import os
import shutil
from io import BytesIO

# Document parsing imports
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import modules
from rolecolor_framework import ROLECOLOR_KEYWORDS, get_rolecolor_descriptions
from resume_scorer import score_resume, format_score_output, get_score_summary
from resume_rewriter import rewrite_summary, enhance_with_chat
from latex_generator import generate_latex_from_resume, generate_simple_latex, edit_latex_with_ai, validate_latex_syntax, generate_tailored_resume_latex
from e2b_runner import compile_latex_in_sandbox, validate_latex_syntax as validate_latex
from chat_assistant import ChatAssistant

# Page configuration
st.set_page_config(
    page_title="RoleColorAI - Resume Intelligence",
    page_icon="🎨",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .rolecolor-builder { color: #3B82F6; font-weight: bold; }
    .rolecolor-enabler { color: #22C55E; font-weight: bold; }
    .rolecolor-thriver { color: #F97316; font-weight: bold; }
    .rolecolor-supportee { color: #8B5CF6; font-weight: bold; }
    
    .score-bar {
        height: 24px;
        border-radius: 4px;
        margin: 4px 0;
    }
    
    .stTextArea textarea {
        font-family: 'Courier New', monospace;
    }
    
    .chat-message {
        padding: 10px;
        border-radius: 8px;
        margin: 5px 0;
    }
    
    .user-message {
        background-color: #E3F2FD;
    }
    
    .assistant-message {
        background-color: #F5F5F5;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = ""
if 'scores' not in st.session_state:
    st.session_state.scores = None
if 'rewritten_summary' not in st.session_state:
    st.session_state.rewritten_summary = ""
if 'latex_code' not in st.session_state:
    st.session_state.latex_code = ""
if 'chat_assistant' not in st.session_state:
    st.session_state.chat_assistant = ChatAssistant()
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = 0
if 'compiled_pdf' not in st.session_state:
    st.session_state.compiled_pdf = None
if 'word_doc' not in st.session_state:
    st.session_state.word_doc = None
if 'job_description' not in st.session_state:
    st.session_state.job_description = ""
if 'tailored_latex' not in st.session_state:
    st.session_state.tailored_latex = ""
if 'tailored_pdf' not in st.session_state:
    st.session_state.tailored_pdf = None
if 'tailored_word' not in st.session_state:
    st.session_state.tailored_word = None


def extract_text_from_pdf(pdf_file) -> str:
    """Extract text from a PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_parts = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""


def extract_text_from_docx(docx_file) -> str:
    """Extract text from a Word document (.docx)."""
    try:
        doc = Document(docx_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return ""


def compile_latex_locally(latex_code: str, filename: str = "resume") -> dict:
    """
    Compile LaTeX to PDF using local pdflatex if available.
    Falls back to an error message if pdflatex is not installed.
    """
    # Check if pdflatex is available
    pdflatex_path = shutil.which("pdflatex")
    
    if not pdflatex_path:
        return {
            "success": False,
            "error": "pdflatex not found. Install TeX Live or MiKTeX for local compilation, or use the E2B option.",
            "log": "",
            "pdf_data": None
        }
    
    try:
        # Create a temporary directory
        with tempfile.TemporaryDirectory() as tmpdir:
            tex_path = os.path.join(tmpdir, f"{filename}.tex")
            pdf_path = os.path.join(tmpdir, f"{filename}.pdf")
            
            # Write LaTeX file
            with open(tex_path, "w", encoding="utf-8") as f:
                f.write(latex_code)
            
            # Run pdflatex
            result = subprocess.run(
                [pdflatex_path, "-interaction=nonstopmode", f"{filename}.tex"],
                cwd=tmpdir,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            # Check if PDF was created
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    pdf_data = base64.b64encode(f.read()).decode("utf-8")
                
                return {
                    "success": True,
                    "pdf_data": pdf_data,
                    "log": result.stdout,
                    "error": None
                }
            else:
                return {
                    "success": False,
                    "error": f"PDF not generated. LaTeX errors may have occurred.",
                    "log": result.stdout + "\n" + result.stderr,
                    "pdf_data": None
                }
    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "error": "Compilation timed out (60s limit)",
            "log": "",
            "pdf_data": None
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "log": "",
            "pdf_data": None
        }


def generate_word_document(resume_text: str, summary: str, dominant_role: str) -> bytes:
    """
    Generate a Word document from resume content.
    Returns the document as bytes.
    """
    doc = Document()
    
    # Define role colors (as RGB tuples)
    role_colors = {
        "Builder": (59, 130, 246),     # Blue
        "Enabler": (34, 197, 94),      # Green
        "Thriver": (249, 115, 22),     # Orange
        "Supportee": (139, 92, 246)    # Purple
    }
    
    # Title
    title = doc.add_heading("Resume", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # RoleColor badge
    role_para = doc.add_paragraph()
    role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_run = role_para.add_run(f"Dominant RoleColor: {dominant_role}")
    role_run.bold = True
    role_run.font.size = Pt(12)
    
    doc.add_paragraph()  # Spacer
    
    # Summary section
    doc.add_heading("Professional Summary", level=1)
    summary_para = doc.add_paragraph(summary)
    summary_para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()  # Spacer
    
    # Parse and add the resume content
    doc.add_heading("Experience & Details", level=1)
    
    # Split resume text into sections
    lines = resume_text.strip().split("\n")
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if it's a section header (all caps or ends with colon)
        if line.isupper() or (len(line) < 50 and line.endswith(":")):
            current_section = line.rstrip(":")
            doc.add_heading(current_section, level=2)
        elif line.startswith("-") or line.startswith("•"):
            # Bullet point
            doc.add_paragraph(line[1:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    
    # Save to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes.getvalue()


def display_rolecolor_scores(scores: dict, matches: dict):
    """Display RoleColor scores with visual bars."""
    colors = {
        "Builder": "#3B82F6",
        "Enabler": "#22C55E", 
        "Thriver": "#F97316",
        "Supportee": "#8B5CF6"
    }
    
    sorted_scores = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    
    for role, score in sorted_scores:
        col1, col2 = st.columns([3, 1])
        with col1:
            st.markdown(f"**{role}**")
            st.progress(score, text=f"{score:.1%}")
        with col2:
            keywords = matches.get(role, [])
            if keywords:
                top_kw = [kw[0] for kw in sorted(keywords, key=lambda x: x[1]*x[2], reverse=True)[:3]]
                st.caption(", ".join(top_kw))


def main():
    # Sidebar - Chat Assistant
    with st.sidebar:
        st.header("💬 AI Chat Assistant")
        st.markdown("Ask questions about your analysis or get help refining your resume.")
        
        # Chat input
        chat_input = st.text_input("Your message:", key="chat_input", placeholder="Ask me anything...")
        
        col1, col2 = st.columns(2)
        with col1:
            send_button = st.button("Send", use_container_width=True)
        with col2:
            clear_button = st.button("Clear Chat", use_container_width=True)
        
        if clear_button:
            st.session_state.chat_history = []
            st.session_state.chat_assistant.clear_history()
            st.rerun()
        
        if send_button and chat_input:
            # Update assistant context
            st.session_state.chat_assistant.set_context(
                resume_text=st.session_state.resume_text,
                scores=st.session_state.scores['scores'] if st.session_state.scores else None,
                summary=st.session_state.rewritten_summary,
                latex_code=st.session_state.latex_code
            )
            
            # Get response
            response = st.session_state.chat_assistant.chat(chat_input)
            
            # Add to history
            st.session_state.chat_history.append({"role": "user", "content": chat_input})
            st.session_state.chat_history.append({"role": "assistant", "content": response})
            
            st.rerun()
        
        # Display chat history
        st.markdown("---")
        chat_container = st.container(height=400)
        with chat_container:
            for msg in st.session_state.chat_history:
                if msg["role"] == "user":
                    st.markdown(f"**You:** {msg['content']}")
                else:
                    st.markdown(f"**AI:** {msg['content']}")
                st.markdown("")
        
        # Quick actions
        st.markdown("---")
        st.subheader("Quick Actions")
        
        if st.button("📊 Explain My Scores", use_container_width=True):
            if st.session_state.scores:
                st.session_state.chat_assistant.set_context(
                    scores=st.session_state.scores['scores']
                )
                response = st.session_state.chat_assistant.explain_scores()
                st.session_state.chat_history.append({"role": "assistant", "content": response})
                st.rerun()
            else:
                st.warning("Analyze a resume first!")
        
        if st.button("✨ Improve Summary", use_container_width=True):
            if st.session_state.rewritten_summary:
                st.session_state.chat_assistant.set_context(
                    summary=st.session_state.rewritten_summary,
                    scores=st.session_state.scores['scores'] if st.session_state.scores else None
                )
                response = st.session_state.chat_assistant.chat(
                    "Suggest improvements for my current resume summary"
                )
                st.session_state.chat_history.append({"role": "assistant", "content": response})
                st.rerun()
            else:
                st.warning("Generate a summary first!")

    # Main content
    st.title("🎨 RoleColorAI")
    st.markdown("*Transform your resume through the lens of team contribution styles*")
    
    # Tab names
    tab_names = [
        "📝 Resume Input", 
        "📊 RoleColor Analysis", 
        "✍️ Summary Rewrite",
        "📄 LaTeX/Overleaf",
        "🎯 Job Tailoring"
    ]
    
    # Tabs for different sections - using session state for active tab
    tab1, tab2, tab3, tab4, tab5 = st.tabs(tab_names)
    
    # Tab 1: Resume Input
    with tab1:
        st.header("Enter Your Resume")
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            # File upload section
            st.subheader("📁 Upload Resume File")
            uploaded_file = st.file_uploader(
                "Upload your resume (PDF or Word)",
                type=["pdf", "docx", "doc"],
                help="Supported formats: PDF (.pdf), Word (.docx)"
            )
            
            if uploaded_file is not None:
                file_type = uploaded_file.name.split(".")[-1].lower()
                
                with st.spinner(f"Extracting text from {uploaded_file.name}..."):
                    if file_type == "pdf":
                        extracted_text = extract_text_from_pdf(uploaded_file)
                    elif file_type in ["docx", "doc"]:
                        extracted_text = extract_text_from_docx(uploaded_file)
                    else:
                        extracted_text = ""
                        st.error("Unsupported file format")
                
                if extracted_text:
                    st.success(f"✅ Successfully extracted text from **{uploaded_file.name}**")
                    st.session_state.resume_text = extracted_text
                    st.info(f"📄 Extracted {len(extracted_text)} characters from your resume")
            
            st.markdown("---")
            st.subheader("📝 Or Paste Text Directly")
            
            resume_input = st.text_area(
                "Paste your resume text here:",
                height=350,
                value=st.session_state.resume_text,
                placeholder="Paste your resume content here...\n\nExample:\nSoftware Engineer with 5 years of experience..."
            )
            
            col_btn1, col_btn2 = st.columns(2)
            with col_btn1:
                if st.button("📤 Load Sample Resume", use_container_width=True):
                    sample = """Senior Software Engineer with 5 years of experience in building scalable 
backend systems and APIs. Led a team of 4 engineers to architect and develop 
a microservices platform that handles 10M+ daily transactions.

EXPERIENCE

Tech Company Inc. - Senior Software Engineer (2021-Present)
- Designed and implemented a real-time data processing pipeline using Kafka and Spark
- Collaborated with cross-functional teams including product, design, and data science
- Mentored 3 junior developers and established coding standards and best practices
- Spearheaded the migration to cloud infrastructure, reducing costs by 40%

Startup Co. - Software Engineer (2019-2021)
- Built REST APIs serving 1M+ daily requests with 99.9% uptime
- Maintained comprehensive documentation for all system components
- Worked in fast-paced agile environment, consistently meeting tight deadlines
- Proactively identified and resolved production issues

SKILLS
Python, Go, AWS, Kubernetes, PostgreSQL, Redis, Kafka, Docker

EDUCATION
B.S. Computer Science, State University, 2019"""
                    st.session_state.resume_text = sample
                    st.rerun()
            
            with col_btn2:
                if st.button("🗑️ Clear Text", use_container_width=True):
                    st.session_state.resume_text = ""
                    st.rerun()
        
        with col2:
            st.subheader("RoleColor Guide")
            
            descriptions = get_rolecolor_descriptions()
            
            with st.expander("🔵 Builder", expanded=False):
                st.markdown(descriptions["Builder"])
            
            with st.expander("🟢 Enabler", expanded=False):
                st.markdown(descriptions["Enabler"])
            
            with st.expander("🟠 Thriver", expanded=False):
                st.markdown(descriptions["Thriver"])
            
            with st.expander("🟣 Supportee", expanded=False):
                st.markdown(descriptions["Supportee"])
        
        if resume_input:
            st.session_state.resume_text = resume_input
            
            if st.button("🔍 Analyze Resume", type="primary", use_container_width=True):
                with st.spinner("Analyzing resume..."):
                    st.session_state.scores = score_resume(resume_input)
                st.success("✅ Analysis complete! Switching to RoleColor Analysis tab...")
                # Auto-switch to RoleColor Analysis tab using JavaScript
                st.markdown("""
                    <script>
                        // Find and click the RoleColor Analysis tab
                        setTimeout(function() {
                            const tabs = window.parent.document.querySelectorAll('[data-baseweb="tab"]');
                            if (tabs.length > 1) {
                                tabs[1].click();
                            }
                        }, 500);
                    </script>
                """, unsafe_allow_html=True)
                # Rerun to reflect changes
                import time
                time.sleep(0.5)
                st.rerun()

    # Tab 2: RoleColor Analysis
    with tab2:
        st.header("RoleColor Analysis Results")
        
        if st.session_state.scores:
            scores = st.session_state.scores
            
            # Dominant RoleColor highlight
            dominant = scores['dominant_role']
            role_colors = {
                "Builder": "🔵", "Enabler": "🟢", 
                "Thriver": "🟠", "Supportee": "🟣"
            }
            
            st.markdown(f"### {role_colors[dominant]} Dominant RoleColor: **{dominant}**")
            st.markdown(get_rolecolor_descriptions()[dominant])
            
            st.markdown("---")
            
            # Score distribution
            col1, col2 = st.columns([2, 1])
            
            with col1:
                st.subheader("Score Distribution")
                display_rolecolor_scores(scores['scores'], scores['matches'])
            
            with col2:
                st.subheader("Summary")
                st.metric("Keywords Matched", scores['total_keywords_matched'])
                
                # Top keywords overall
                all_keywords = []
                for role, matches in scores['matches'].items():
                    for kw, weight, count in matches:
                        all_keywords.append((kw, role, weight * count))
                
                all_keywords.sort(key=lambda x: x[2], reverse=True)
                
                st.markdown("**Top Keywords Found:**")
                for kw, role, _ in all_keywords[:8]:
                    st.markdown(f"• {kw} ({role})")
            
            # Raw scores expander
            with st.expander("📈 Detailed Scoring Data"):
                st.json({
                    "normalized_scores": {k: round(v, 4) for k, v in scores['scores'].items()},
                    "raw_scores": {k: round(v, 2) for k, v in scores['raw_scores'].items()},
                    "keyword_counts": {k: len(v) for k, v in scores['matches'].items()}
                })
        else:
            st.info("👈 Enter a resume in the Resume Input tab and click Analyze")

    # Tab 3: Summary Rewrite
    with tab3:
        st.header("AI-Powered Summary Rewrite")
        
        if st.session_state.scores:
            dominant = st.session_state.scores['dominant_role']
            
            st.markdown(f"Generating summary optimized for **{dominant}** traits.")
            
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("✨ Generate Summary", type="primary", use_container_width=True):
                    with st.spinner("AI is writing your summary..."):
                        summary = rewrite_summary(
                            resume_text=st.session_state.resume_text,
                            dominant_role=dominant,
                            scores=st.session_state.scores['scores']
                        )
                        st.session_state.rewritten_summary = summary
            
            with col2:
                # Get the dominant role and set it as default selection
                role_options = ["Builder", "Enabler", "Thriver", "Supportee"]
                dominant_index = role_options.index(dominant) if dominant in role_options else 0
                
                target_role = st.selectbox(
                    "Target a specific RoleColor:",
                    role_options,
                    index=dominant_index,
                    help=f"Auto-selected based on your dominant RoleColor: {dominant}"
                )
                
                if st.button("🎯 Generate for Selected", use_container_width=True):
                    with st.spinner("AI is writing your summary..."):
                        summary = rewrite_summary(
                            resume_text=st.session_state.resume_text,
                            dominant_role=target_role,
                            scores=st.session_state.scores['scores']
                        )
                        st.session_state.rewritten_summary = summary
                        st.rerun()
            
            if st.session_state.rewritten_summary:
                st.markdown("---")
                st.subheader("Generated Summary")
                
                # Editable summary
                edited_summary = st.text_area(
                    "Edit your summary:",
                    value=st.session_state.rewritten_summary,
                    height=200
                )
                
                if edited_summary != st.session_state.rewritten_summary:
                    st.session_state.rewritten_summary = edited_summary
                
                # Refinement through chat
                st.markdown("---")
                st.subheader("Refine with AI")
                
                refinement = st.text_input(
                    "What would you like to change?",
                    placeholder="e.g., Make it more technical, Add leadership focus, Shorter..."
                )
                
                if refinement and st.button("Apply Changes"):
                    with st.spinner("Refining..."):
                        new_summary = enhance_with_chat(
                            resume_text=st.session_state.resume_text,
                            current_summary=st.session_state.rewritten_summary,
                            user_feedback=refinement,
                            dominant_role=dominant
                        )
                        st.session_state.rewritten_summary = new_summary
                        st.rerun()
        else:
            st.info("👈 Analyze a resume first to generate a summary")

    # Tab 4: LaTeX/Overleaf
    with tab4:
        st.header("LaTeX Resume Generator")
        
        if st.session_state.scores and st.session_state.rewritten_summary:
            dominant = st.session_state.scores['dominant_role']
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.subheader("Generate LaTeX")
                
                generation_type = st.radio(
                    "Generation mode:",
                    ["Full Resume", "Summary Only"],
                    horizontal=True
                )
                
                if st.button("📄 Generate LaTeX Code", type="primary", use_container_width=True):
                    with st.spinner("Generating LaTeX..."):
                        if generation_type == "Full Resume":
                            latex = generate_latex_from_resume(
                                resume_text=st.session_state.resume_text,
                                rewritten_summary=st.session_state.rewritten_summary,
                                dominant_role=dominant,
                                scores=st.session_state.scores['scores']
                            )
                        else:
                            latex = generate_simple_latex(
                                name="Your Name",
                                contact="email@example.com | (555) 123-4567",
                                summary=st.session_state.rewritten_summary,
                                dominant_role=dominant
                            )
                        st.session_state.latex_code = latex
                        st.session_state.compiled_pdf = None  # Reset PDF when new LaTeX is generated
                
                # Validate syntax
                if st.session_state.latex_code:
                    validation = validate_latex(st.session_state.latex_code)
                    if validation['valid']:
                        st.success("✅ LaTeX syntax is valid")
                    else:
                        for error in validation['errors']:
                            st.error(f"❌ {error}")
                        for warning in validation['warnings']:
                            st.warning(f"⚠️ {warning}")
            
            with col2:
                st.subheader("📥 Download Options")
                
                if st.session_state.latex_code:
                    # LaTeX source download
                    st.download_button(
                        label="📄 Download LaTeX (.tex)",
                        data=st.session_state.latex_code,
                        file_name="resume.tex",
                        mime="text/plain",
                        use_container_width=True
                    )
                    
                    # Word document download
                    st.markdown("---")
                    st.markdown("**Word Document:**")
                    if st.button("📝 Generate Word Document", use_container_width=True):
                        with st.spinner("Generating Word document..."):
                            word_bytes = generate_word_document(
                                resume_text=st.session_state.resume_text,
                                summary=st.session_state.rewritten_summary,
                                dominant_role=dominant
                            )
                            st.session_state.word_doc = word_bytes
                    
                    if 'word_doc' in st.session_state and st.session_state.word_doc:
                        st.download_button(
                            label="📥 Download Word (.docx)",
                            data=st.session_state.word_doc,
                            file_name="resume.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                else:
                    st.info("Generate LaTeX code first")
            
            # PDF Compilation Section
            st.markdown("---")
            st.subheader("🖨️ PDF Compilation")
            
            if st.session_state.latex_code:
                comp_col1, comp_col2 = st.columns(2)
                
                with comp_col1:
                    st.markdown("**Option 1: Local Compilation**")
                    st.caption("Uses pdflatex if installed on your system")
                    
                    if st.button("🖥️ Compile Locally", use_container_width=True):
                        with st.spinner("Compiling LaTeX locally..."):
                            result = compile_latex_locally(st.session_state.latex_code)
                            
                            if result['success']:
                                st.success("✅ PDF compiled successfully!")
                                st.session_state.compiled_pdf = result['pdf_data']
                            else:
                                st.error(f"❌ {result['error']}")
                                if result['log']:
                                    with st.expander("View compilation log"):
                                        st.code(result['log'])
                
                with comp_col2:
                    st.markdown("**Option 2: E2B Cloud Compilation**")
                    st.caption("Uses E2B sandbox (requires API key)")
                    
                    if st.button("☁️ Compile in E2B", use_container_width=True):
                        with st.spinner("Compiling LaTeX in E2B sandbox... (this may take a minute)"):
                            result = compile_latex_in_sandbox(st.session_state.latex_code)
                            
                            if result['success']:
                                st.success("✅ PDF compiled successfully!")
                                st.session_state.compiled_pdf = result['pdf_data']
                            else:
                                st.error(f"❌ Compilation failed: {result['error']}")
                                if result['log']:
                                    with st.expander("View compilation log"):
                                        st.code(result['log'])
                
                # Show PDF download if compiled
                if st.session_state.compiled_pdf:
                    st.markdown("---")
                    pdf_bytes = base64.b64decode(st.session_state.compiled_pdf)
                    st.download_button(
                        label="📥 Download PDF",
                        data=pdf_bytes,
                        file_name="resume.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        type="primary"
                    )
                    st.success("✅ Your PDF is ready for download!")
            
            # LaTeX Editor
            st.markdown("---")
            st.subheader("✏️ LaTeX Editor")
            
            if st.session_state.latex_code:
                # AI modification input
                latex_modification = st.text_input(
                    "Ask AI to modify the LaTeX:",
                    placeholder="e.g., Add a projects section, Change the color scheme, Make margins wider..."
                )
                
                if latex_modification and st.button("🤖 Apply AI Modification"):
                    with st.spinner("Modifying LaTeX..."):
                        new_latex = edit_latex_with_ai(
                            st.session_state.latex_code,
                            latex_modification
                        )
                        st.session_state.latex_code = new_latex
                        st.session_state.compiled_pdf = None  # Reset PDF after modification
                        st.rerun()
                
                # Code editor
                edited_latex = st.text_area(
                    "Edit LaTeX code directly:",
                    value=st.session_state.latex_code,
                    height=400
                )
                
                if edited_latex != st.session_state.latex_code:
                    st.session_state.latex_code = edited_latex
                    st.session_state.compiled_pdf = None  # Reset PDF after edit
                
                # Overleaf tip
                st.info("💡 **Tip:** Copy the LaTeX code above and paste it into [Overleaf](https://www.overleaf.com/) for online editing and compilation if local/E2B compilation doesn't work.")
        else:
            st.info("👈 Complete the analysis and generate a summary first")

    # Tab 5: Job Tailoring
    with tab5:
        st.header("🎯 Job-Tailored Resume Generator")
        st.markdown("*Generate a resume tailored specifically to a job description*")
        
        if st.session_state.resume_text:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.subheader("📋 Job Description")
                job_desc_input = st.text_area(
                    "Paste the job description here:",
                    value=st.session_state.job_description,
                    height=300,
                    placeholder="Paste the full job posting here...\n\nExample:\nSenior Software Engineer\nWe are looking for an experienced developer..."
                )
                
                if job_desc_input != st.session_state.job_description:
                    st.session_state.job_description = job_desc_input
                
                # Get dominant role for styling (if analysis done)
                dominant_role = "Builder"  # Default
                if st.session_state.scores:
                    dominant_role = st.session_state.scores['dominant_role']
                
                st.markdown("---")
                st.markdown("**Resume Style:**")
                style_role = st.selectbox(
                    "Color accent based on RoleColor:",
                    ["Builder", "Enabler", "Thriver", "Supportee"],
                    index=["Builder", "Enabler", "Thriver", "Supportee"].index(dominant_role),
                    help="This affects the color scheme of your tailored resume"
                )
            
            with col2:
                st.subheader("📄 Your Current Resume")
                st.text_area(
                    "Resume preview:",
                    value=st.session_state.resume_text[:1500] + "..." if len(st.session_state.resume_text) > 1500 else st.session_state.resume_text,
                    height=300,
                    disabled=True
                )
            
            # Generate button
            st.markdown("---")
            if st.button("✨ Generate Tailored Resume", type="primary", use_container_width=True):
                if not st.session_state.job_description.strip():
                    st.error("⚠️ Please paste a job description first!")
                else:
                    with st.spinner("🤖 AI is analyzing the job and tailoring your resume... This may take a moment."):
                        tailored = generate_tailored_resume_latex(
                            resume_text=st.session_state.resume_text,
                            job_description=st.session_state.job_description,
                            dominant_role=style_role
                        )
                        st.session_state.tailored_latex = tailored
                        st.session_state.tailored_pdf = None  # Reset compiled PDF
                        st.session_state.tailored_word = None  # Reset Word doc
                    st.success("✅ Tailored resume generated!")
                    st.rerun()
            
            # Show generated LaTeX and options
            if st.session_state.tailored_latex:
                st.markdown("---")
                st.subheader("📝 Generated Tailored Resume")
                
                # Validation
                validation = validate_latex(st.session_state.tailored_latex)
                if validation['valid']:
                    st.success("✅ LaTeX syntax is valid")
                else:
                    for error in validation['errors']:
                        st.warning(f"⚠️ {error}")
                
                # Options columns
                opt_col1, opt_col2, opt_col3 = st.columns(3)
                
                with opt_col1:
                    st.markdown("**📥 Download Options**")
                    st.download_button(
                        label="📄 Download LaTeX (.tex)",
                        data=st.session_state.tailored_latex,
                        file_name="tailored_resume.tex",
                        mime="text/plain",
                        use_container_width=True
                    )
                    
                    # Word document generation
                    if st.button("📝 Generate Word Doc", key="gen_tailored_word", use_container_width=True):
                        with st.spinner("Generating Word document..."):
                            # Extract a summary from the tailored resume for the Word doc
                            summary = "Professional with experience aligned to your target role."
                            if st.session_state.rewritten_summary:
                                summary = st.session_state.rewritten_summary
                            
                            word_bytes = generate_word_document(
                                resume_text=st.session_state.resume_text,
                                summary=summary,
                                dominant_role=style_role
                            )
                            st.session_state.tailored_word = word_bytes
                    
                    if st.session_state.tailored_word:
                        st.download_button(
                            label="📥 Download Word (.docx)",
                            data=st.session_state.tailored_word,
                            file_name="tailored_resume.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                
                with opt_col2:
                    st.markdown("**🖨️ Local Compilation**")
                    if st.button("🖥️ Compile Locally", key="local_tailored", use_container_width=True):
                        with st.spinner("Compiling LaTeX locally..."):
                            result = compile_latex_locally(st.session_state.tailored_latex)
                            
                            if result['success']:
                                st.success("✅ PDF compiled!")
                                st.session_state.tailored_pdf = result['pdf_data']
                            else:
                                st.error(f"❌ {result['error']}")
                
                with opt_col3:
                    st.markdown("**☁️ E2B Compilation**")
                    if st.button("☁️ Compile in E2B", key="e2b_tailored", use_container_width=True):
                        with st.spinner("Compiling in E2B sandbox..."):
                            result = compile_latex_in_sandbox(st.session_state.tailored_latex)
                            
                            if result['success']:
                                st.success("✅ PDF compiled!")
                                st.session_state.tailored_pdf = result['pdf_data']
                            else:
                                st.error(f"❌ {result['error']}")
                
                # PDF download (if compiled)
                if st.session_state.tailored_pdf:
                    st.markdown("---")
                    pdf_bytes = base64.b64decode(st.session_state.tailored_pdf)
                    st.download_button(
                        label="📥 Download Tailored PDF",
                        data=pdf_bytes,
                        file_name="tailored_resume.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        type="primary"
                    )
                    st.success("🎉 Your tailored PDF is ready!")
                
                # LaTeX editor
                st.markdown("---")
                st.subheader("✏️ Edit Tailored LaTeX")
                
                edit_request = st.text_input(
                    "Ask AI to modify:",
                    placeholder="e.g., Emphasize leadership more, Add a projects section, Make it more concise...",
                    key="tailored_edit_request"
                )
                
                if edit_request and st.button("🤖 Apply Modification", key="apply_tailored_edit"):
                    with st.spinner("Modifying LaTeX..."):
                        new_latex = edit_latex_with_ai(
                            st.session_state.tailored_latex,
                            edit_request
                        )
                        st.session_state.tailored_latex = new_latex
                        st.session_state.tailored_pdf = None
                        st.rerun()
                
                # Manual editor
                edited_tailored = st.text_area(
                    "Edit LaTeX directly:",
                    value=st.session_state.tailored_latex,
                    height=400,
                    key="tailored_latex_editor"
                )
                
                if edited_tailored != st.session_state.tailored_latex:
                    st.session_state.tailored_latex = edited_tailored
                    st.session_state.tailored_pdf = None
                
                st.info("💡 **Tip:** Copy this LaTeX into [Overleaf](https://www.overleaf.com/) for online editing!")
        else:
            st.info("👈 Please enter your resume in the **Resume Input** tab first to use Job Tailoring.")
            st.markdown("""
            ### How to use Job Tailoring:
            1. **Go to Resume Input tab** and paste or upload your resume
            2. **Come back here** and paste the target job description
            3. **Click Generate** to create a tailored resume
            4. **Download** as LaTeX, PDF, or Word document
            """)

    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: gray;'>
            RoleColorAI - Resume Intelligence System<br>
            Analyzing team contribution styles through AI
        </div>
        """,
        unsafe_allow_html=True
    )


if __name__ == "__main__":
    main()
